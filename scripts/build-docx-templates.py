#!/usr/bin/env python3
"""Build the Arabic operational Word template library.

Generated documents are intentionally blank institutional forms. Bracketed values
are prompts, not production data. The output is suitable for direct download from
the committee portal after visual and privacy QA.
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("public/templates")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

GREEN = "17685D"
DEEP_GREEN = "103E38"
MINT = "EAF3F0"
GOLD = "B58A45"
GOLD_LIGHT = "F6EEDC"
INK = "17312D"
MUTED = "64756F"
LINE = "D9E3DF"
PALE = "F5F8F6"
WHITE = "FFFFFF"
RED_PALE = "F9EAEA"
FONT = "Arial"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_rtl(paragraph, *, alignment=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    paragraph.alignment = alignment
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_run_font(run, *, size: float = 10.5, bold: bool = False, color: str = INK) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def set_document_defaults(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12
    for style_name, size, color in (("Title", 25, DEEP_GREEN), ("Heading 1", 17, DEEP_GREEN), ("Heading 2", 13, GREEN), ("Heading 3", 11.5, INK)):
        style = styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    props = doc.core_properties
    props.title = title
    props.subject = "قالب تشغيلي لأعمال اللجان"
    props.author = "بوابة أعمال اللجان"
    props.last_modified_by = "بوابة أعمال اللجان"
    props.keywords = "لجان، جودة، دليل، اعتماد"
    props.comments = "قالب فارغ؛ الحقول بين معقوفين تُستكمل قبل الاعتماد."

    def populate_header(header) -> None:
        header_table = header.add_table(rows=1, cols=2, width=Cm(17.6))
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_table.autofit = False
        header_table.columns[0].width = Cm(11.8)
        header_table.columns[1].width = Cm(5.8)
        for cell in header_table.rows[0].cells:
            set_cell_shading(cell, DEEP_GREEN)
            set_cell_border(cell, DEEP_GREEN)
            set_cell_margins(cell, 70, 130, 70, 130)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = header_table.cell(0, 1).paragraphs[0]
        set_paragraph_rtl(p)
        set_run_font(p.add_run("كلية الشريعة والأنظمة"), size=9.5, bold=True, color=WHITE)
        p = header_table.cell(0, 0).paragraphs[0]
        set_paragraph_rtl(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_run_font(p.add_run(title), size=8.7, bold=True, color=WHITE)

    def populate_footer(footer) -> None:
        table = footer.add_table(rows=1, cols=2, width=Cm(17.6))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(12.0)
        table.columns[1].width = Cm(5.6)
        for cell in table.rows[0].cells:
            set_cell_shading(cell, PALE)
            set_cell_border(cell, LINE)
            set_cell_margins(cell, 60, 110, 60, 110)
        p = table.cell(0, 1).paragraphs[0]
        set_paragraph_rtl(p)
        set_run_font(p.add_run("إصدار 2.0 · قالب غير مكتمل حتى الاعتماد"), size=8, color=MUTED)
        p = table.cell(0, 0).paragraphs[0]
        set_paragraph_rtl(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_run_font(p.add_run("صفحة "), size=8, color=MUTED)
        add_page_field(p)

    doc.settings.odd_and_even_pages_header_footer = True
    populate_header(section.header)
    populate_header(section.even_page_header)
    populate_footer(section.footer)
    populate_footer(section.even_page_footer)


def add_text(doc: Document, text: str, *, size: float = 10.5, bold: bool = False, color: str = INK,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT, before: float = 0, after: float = 4) -> None:
    p = doc.add_paragraph()
    set_paragraph_rtl(p, alignment=alignment)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    set_run_font(p.add_run(text), size=size, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_rtl(p)
    set_run_font(p.add_run(text), size={1: 17, 2: 13, 3: 11.5}[level], bold=True, color=DEEP_GREEN if level == 1 else GREEN)


def add_callout(doc: Document, title: str, body: str, *, tone: str = "green") -> None:
    fill, accent = (MINT, GREEN) if tone == "green" else (GOLD_LIGHT, GOLD) if tone == "gold" else (RED_PALE, "A64C4C")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17.4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent, "8")
    set_cell_margins(cell, 120, 150, 120, 150)
    p = cell.paragraphs[0]
    set_paragraph_rtl(p)
    set_run_font(p.add_run(f"{title}: "), size=10, bold=True, color=accent)
    set_run_font(p.add_run(body), size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_filename_callout(doc: Document, filename: str) -> None:
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(17.4)
    label_cell, value_cell = table.column_cells(0)
    for cell in (label_cell, value_cell):
        set_cell_shading(cell, MINT)
        set_cell_border(cell, GREEN, "8")
        set_cell_margins(cell, 85, 140, 85, 140)
    p = label_cell.paragraphs[0]
    set_paragraph_rtl(p)
    set_run_font(p.add_run("اسم الملف المقترح"), size=9.5, bold=True, color=GREEN)
    p = value_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.add_run(filename), size=9.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_fields_table(doc: Document, rows: Iterable[tuple[str, str]], *, widths: tuple[float, float] = (5.0, 12.4)) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(widths[0])
    table.columns[1].width = Cm(widths[1])
    for index, (label, value) in enumerate(rows):
        cells = table.add_row().cells
        cells[0].width = Cm(widths[0])
        cells[1].width = Cm(widths[1])
        for cell in cells:
            set_cell_border(cell)
            set_cell_margins(cell, 95, 120, 95, 120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cells[0], MINT if index % 2 == 0 else PALE)
        set_cell_shading(cells[1], WHITE)
        p = cells[0].paragraphs[0]
        set_paragraph_rtl(p)
        set_run_font(p.add_run(label), size=9.4, bold=True, color=GREEN)
        p = cells[1].paragraphs[0]
        set_paragraph_rtl(p)
        set_run_font(p.add_run(value), size=9.4, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_grid(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None,
             font_size: float = 8.6) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths is None:
        widths = [17.4 / len(headers)] * len(headers)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Cm(widths[idx])
        set_cell_shading(cell, DEEP_GREEN)
        set_cell_border(cell, DEEP_GREEN)
        set_cell_margins(cell, 85, 80, 85, 80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph_rtl(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(p.add_run(header), size=font_size, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].width = Cm(widths[idx])
            set_cell_shading(cells[idx], WHITE if row_index % 2 == 0 else PALE)
            set_cell_border(cells[idx])
            set_cell_margins(cells[idx], 90, 80, 90, 80)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            set_paragraph_rtl(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT if idx else WD_ALIGN_PARAGRAPH.CENTER)
            set_run_font(p.add_run(value), size=font_size, color=MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_checklist(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        set_paragraph_rtl(p)
        p.paragraph_format.left_indent = Cm(0.2)
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run("☐  "), size=11, bold=True, color=GREEN)
        set_run_font(p.add_run(item), size=9.6, color=INK)


def add_numbered_steps(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, 1):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Cm(1.2)
        table.columns[1].width = Cm(16.2)
        number_cell, text_cell = table.rows[0].cells
        set_cell_shading(number_cell, GREEN)
        set_cell_border(number_cell, GREEN)
        set_cell_shading(text_cell, PALE)
        set_cell_border(text_cell)
        for cell in (number_cell, text_cell):
            set_cell_margins(cell, 85, 90, 85, 90)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = number_cell.paragraphs[0]
        set_paragraph_rtl(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_run_font(p.add_run(str(index)), size=10, bold=True, color=WHITE)
        p = text_cell.paragraphs[0]
        set_paragraph_rtl(p)
        set_run_font(p.add_run(item), size=9.5, color=INK)


def add_cover(doc: Document, title: str, subtitle: str, template_id: str) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(34)
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    set_run_font(p.add_run("بوابة أعمال اللجان"), size=11, bold=True, color=GOLD)
    p = doc.add_paragraph(style="Title")
    set_paragraph_rtl(p)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(title), size=25, bold=True, color=DEEP_GREEN)
    add_text(doc, subtitle, size=12, color=MUTED, after=18)
    add_callout(doc, "قاعدة الاستخدام", "هذا قالب عمل فارغ. جميع العبارات بين [معقوفين] حقول يجب استكمالها، وأي سطر موسوم «مثال توضيحي» لا يمثل بيانات فعلية.", tone="gold")
    add_heading(doc, "هوية الملف", 2)
    add_fields_table(doc, [
        ("رمز القالب", template_id),
        ("رمز المهمة", "[مثل: QRA-T003]"),
        ("القسم / اللجنة", "[اسم القسم] / [اسم اللجنة]"),
        ("الفصل والأسبوع", "[1448-F1] / [W01]"),
        ("إصدار الملف", "v01 · [تاريخ التحديث]"),
        ("حالة الوثيقة", "مسودة / تحت المراجعة / معتمدة"),
    ])
    add_heading(doc, "قاعدة التسمية الموحدة", 2)
    add_filename_callout(doc, "1448-F1_W01_QRA_[COMMITTEE]_[TASK-ID]_[TEMPLATE]_v01.docx")
    add_text(doc, "يحفظ الملف في المسار التنظيمي المقترح داخل SharePoint بعد اعتماد مسؤول المنصة. لا تعد هذه الصفحة إثباتًا على الرفع أو الاعتماد.", size=9.5, color=MUTED)


def add_approval_page(doc: Document, checks: list[str], *, include_data_quality: bool = True) -> None:
    doc.add_page_break()
    add_heading(doc, "المراجعة والاعتماد والإقفال", 1)
    add_callout(doc, "حلقة الجودة", "تخطيط ← تنفيذ ← نتيجة موثقة ← تحليل ← تحسين ← تحقق من الإقفال.", tone="green")
    add_heading(doc, "قائمة الفحص النهائي", 2)
    common = [
        "اكتملت الحقول الإلزامية ولم تبقَ قيم افتراضية أو أمثلة توضيحية.",
        "كل نتيجة أو حكم مرتبط بمصدر أو شاهد يمكن التحقق منه.",
        "لكل توصية مسؤول وموعد وحالة متابعة ودليل إقفال متوقع.",
        "رُوجعت الخصوصية، ولم تُدرج بيانات شخصية غير لازمة أو روابط مفتوحة للعامة.",
        "طابق اسم الملف قاعدة التسمية، وحُفظت نسخة PDF عند الحاجة.",
    ]
    if include_data_quality:
        common.insert(2, "تمت مراجعة الاتساق والأرقام والتواريخ والمعادلات قبل الاعتماد.")
    add_checklist(doc, checks + common)
    add_heading(doc, "سجل الاعتماد", 2)
    add_grid(doc, ["الدور", "الاسم/الصفة", "القرار", "التاريخ", "التوقيع/المرجع"], [
        ["1", "[معد الوثيقة]", "أعدّ", "[تاريخ]", "[مرجع]"],
        ["2", "[مراجع اللجنة]", "راجع / أعاد", "[تاريخ]", "[مرجع]"],
        ["3", "[رئيس اللجنة]", "اعتمد / رفض", "[تاريخ]", "[مرجع]"],
        ["4", "[الجهة الإشرافية عند اللزوم]", "اطلع / وجّه", "[تاريخ]", "[مرجع]"],
    ], widths=[1.0, 4.5, 3.5, 3.1, 5.3], font_size=8.2)
    add_heading(doc, "سجل التحسين والإقفال", 2)
    add_grid(doc, ["#", "الفجوة/الفرصة", "الإجراء", "المسؤول", "الموعد", "حالة التحقق"], [
        [str(i), "[وصف محدد]", "[إجراء قابل للقياس]", "[الدور]", "[تاريخ]", "مفتوح / متحقق"] for i in range(1, 4)
    ], widths=[0.8, 4.0, 4.3, 3.0, 2.7, 2.6], font_size=7.9)
    add_callout(doc, "صلة الدراسة الذاتية", "[المعيار/المحك] · [رقم الدليل] · [وصف سبب صلاحية هذا الملف للاستشهاد]", tone="gold")


def build_minutes() -> Document:
    title = "قالب محضر اجتماع لجنة"
    doc = Document()
    set_document_defaults(doc, title)
    add_cover(doc, title, "محضر يحوّل النقاش إلى قرارات قابلة للتتبع والإقفال.", "minutes")
    doc.add_page_break()
    add_heading(doc, "بيانات الاجتماع والحوكمة", 1)
    add_fields_table(doc, [
        ("اسم اللجنة", "[الاسم الرسمي كما في قرار التشكيل]"), ("رقم الاجتماع", "[رقم متسلسل]"),
        ("التاريخ والوقت", "[هجري] / [ميلادي] · من [وقت] إلى [وقت]"), ("مكان/وسيلة الانعقاد", "[حضوري / افتراضي / هجين]"),
        ("رئيس الاجتماع", "[الصفة دون بيانات اتصال]"), ("أمين المحضر", "[الصفة]"),
        ("مرجع الدعوة", "[رقم أو رابط داخلي مقيد]"), ("النصاب", "[مكتمل / غير مكتمل مع التفسير]"),
    ])
    add_heading(doc, "الحضور والاعتذارات", 2)
    add_grid(doc, ["#", "الاسم/الصفة", "صفة المشاركة", "الحضور", "ملاحظة"], [
        [str(i), "[يُستكمل]", "عضو / مقرر / مدعو خبير", "حاضر / معتذر", "[عند الحاجة]"] for i in range(1, 6)
    ], widths=[0.8, 5.0, 4.0, 3.2, 4.4], font_size=8.1)
    add_callout(doc, "تنبيه حوكمة", "المدعو أو الخبير الخارجي لا يُعد عضوًا، ولا يدخل في النصاب أو التصويت، ولا يُكلّف بمهمة داخلية إلا وفق صلاحية موثقة.", tone="gold")
    doc.add_page_break()
    add_heading(doc, "جدول الأعمال والمناقشات", 1)
    add_grid(doc, ["البند", "الموضوع", "المدخلات/الوثائق", "خلاصة المناقشة", "النتيجة"], [
        [str(i), "[عنوان مختصر]", "[مرجع المرفق]", "[حقائق وآراء رئيسة]", "قرار / إحاطة / تأجيل"] for i in range(1, 5)
    ], widths=[1.0, 3.3, 3.7, 6.0, 3.4], font_size=7.9)
    add_heading(doc, "مصفوفة القرارات والتكليفات", 2)
    add_grid(doc, ["#", "نص القرار", "المسؤول", "الموعد", "دليل الإقفال", "الحالة"], [
        [str(i), "[قرار واضح قابل للتنفيذ]", "[دور/جهة]", "[تاريخ]", "[اسم/رابط داخلي]", "مفتوح"] for i in range(1, 5)
    ], widths=[0.8, 4.7, 3.0, 2.7, 3.8, 2.4], font_size=7.8)
    add_heading(doc, "المرفقات", 2)
    add_grid(doc, ["#", "اسم المرفق", "الغرض", "التصنيف", "مرجع الحفظ"], [
        [str(i), "[اسم موحد]", "[ما يثبته]", "داخلي / مقيد", "[مسار داخلي]"] for i in range(1, 4)
    ], widths=[0.8, 4.9, 4.1, 3.2, 4.4], font_size=8.1)
    add_approval_page(doc, ["طابق نص كل قرار ما اتفق عليه الحضور.", "أُثبتت الاعتذارات ولم يُحسب المدعو ضمن النصاب."])
    return doc


def build_operational_plan() -> Document:
    title = "قالب خطة تشغيلية للجنة"
    doc = Document()
    set_document_defaults(doc, title)
    add_cover(doc, title, "خطة تربط التكليف بالهدف والمؤشر والمستهدف والموارد والمتابعة.", "operational-plan")
    doc.add_page_break()
    add_heading(doc, "السياق ونطاق الخطة", 1)
    add_fields_table(doc, [
        ("مرجع تشكيل اللجنة", "[قرار/تكليف]"), ("مدة الخطة", "[من] إلى [إلى]"),
        ("نطاق المسؤولية", "[ما يدخل وما لا يدخل]"), ("الجهة الإشرافية", "[رئيس القسم / عميد الكلية حسب الصلاحية]"),
        ("أصحاب المصلحة", "[جهات داخلية ذات صلة]"), ("المصادر المرجعية", "[لوائح/تقارير/نتائج سابقة]"),
    ])
    add_heading(doc, "تحليل نقطة البداية", 2)
    add_grid(doc, ["البعد", "الوضع الحالي", "الشاهد", "الفجوة/الفرصة", "الأولوية"], [
        ["1", "[وصف واقعي]", "[مصدر]", "[فجوة محددة]", "عالية/متوسطة/منخفضة"],
        ["2", "[وصف واقعي]", "[مصدر]", "[فرصة محددة]", "عالية/متوسطة/منخفضة"],
        ["3", "[وصف واقعي]", "[مصدر]", "[قيد/مخاطرة]", "عالية/متوسطة/منخفضة"],
    ], widths=[1.0, 4.1, 3.1, 5.7, 3.5], font_size=8.0)
    add_heading(doc, "الأهداف ومؤشرات النجاح", 2)
    add_grid(doc, ["#", "الهدف القابل للقياس", "المؤشر", "خط الأساس", "المستهدف", "مصدر القياس"], [
        [str(i), "[فعل + نتيجة + نطاق]", "[اسم/معادلة المؤشر]", "[قيمة/لا يتوفر]", "[قيمة + تاريخ]", "[مصدر موثوق]"] for i in range(1, 4)
    ], widths=[0.8, 4.6, 3.4, 2.6, 2.7, 3.3], font_size=7.7)
    doc.add_page_break()
    add_heading(doc, "الخطة التنفيذية", 1)
    add_grid(doc, ["#", "الإجراء/المخرج", "المالك", "البداية", "النهاية", "المورد", "دليل الإنجاز"], [
        [str(i), "[إجراء محدد]", "[الدور]", "[تاريخ]", "[تاريخ]", "[مورد]", "[دليل متوقع]"] for i in range(1, 7)
    ], widths=[0.7, 4.1, 2.6, 2.2, 2.2, 2.6, 3.0], font_size=7.35)
    add_heading(doc, "إدارة المخاطر", 2)
    add_grid(doc, ["الخطر", "الاحتمال", "الأثر", "الاستجابة", "المالك", "إشارة الإنذار"], [
        ["[خطر محدد]", "1–5", "1–5", "تجنب/خفض/نقل/قبول", "[الدور]", "[شرط]"] for _ in range(3)
    ], widths=[4.3, 2.1, 2.0, 3.7, 2.5, 2.8], font_size=7.8)
    add_heading(doc, "دورية المتابعة", 2)
    add_fields_table(doc, [("اجتماع المتابعة", "[أسبوعي/شهري]"), ("مالك تحديث المؤشرات", "[الدور]"), ("قناة التصعيد", "رئيس اللجنة ← رئيس القسم ← عميد الكلية حسب الصلاحية"), ("موعد مراجعة الخطة", "[تاريخ]")])
    add_approval_page(doc, ["لا يوجد هدف بلا مؤشر ومستهدف ومصدر قياس.", "لا يوجد إجراء بلا مسؤول وموعد ودليل متوقع."])
    return doc


def build_completion_report() -> Document:
    title = "قالب تقرير إنجاز وتحليل مهمة"
    doc = Document()
    set_document_defaults(doc, title)
    add_cover(doc, title, "تقرير يثبت التنفيذ ويحلل النتيجة ويحوّلها إلى تحسين قابل للإقفال.", "completion-report")
    doc.add_page_break()
    add_heading(doc, "بطاقة المهمة ومنهج التنفيذ", 1)
    add_fields_table(doc, [
        ("عنوان المهمة", "[العنوان المحدد]"), ("رمز المهمة", "[TASK-ID]"),
        ("الهدف", "[النتيجة المراد تحقيقها]"), ("النطاق", "[المشمول/غير المشمول]"),
        ("المسؤول", "[الدور]"), ("فترة التنفيذ", "[من] إلى [إلى]"),
        ("المدخلات", "[بيانات/لوائح/نتائج سابقة]"), ("طريقة التحقق", "[كيف نعرف أن التنفيذ صحيح؟]"),
    ])
    add_heading(doc, "خطوات التنفيذ الفعلية", 2)
    add_grid(doc, ["#", "الخطوة", "ما تم فعليًا", "المخرج الوسيط", "التاريخ", "الشاهد"], [
        [str(i), "[خطوة مخططة]", "[وصف موجز]", "[مخرج]", "[تاريخ]", "[مرجع]"] for i in range(1, 5)
    ], widths=[0.8, 3.5, 4.8, 3.3, 2.3, 2.7], font_size=7.8)
    add_callout(doc, "مثال توضيحي", "لا تكتب «تم التنفيذ بنجاح» وحدها. اذكر حجم التنفيذ ونتيجته والشاهد الذي يسمح للمراجع بإعادة التحقق.", tone="gold")
    doc.add_page_break()
    add_heading(doc, "النتائج والتحليل", 1)
    add_grid(doc, ["المؤشر/السؤال", "المستهدف", "النتيجة", "الفرق", "المصدر", "الحكم"], [
        ["[مؤشر]", "[قيمة]", "[قيمة]", "[± قيمة]", "[مصدر]", "تحقق/جزئي/لم يتحقق"] for _ in range(4)
    ], widths=[3.8, 2.4, 2.4, 2.2, 3.1, 3.5], font_size=7.8)
    add_heading(doc, "تفسير النتائج", 2)
    add_fields_table(doc, [("ماذا حدث؟", "[أهم نتيجة مدعومة بالبيانات]"), ("لماذا حدث؟", "[أسباب محتملة ومدى قوة الدليل]"), ("ما أثره؟", "[أثر تشغيلي/أكاديمي]"), ("ما حدود الاستنتاج؟", "[نقص بيانات/تحيز/فترة قصيرة]")])
    add_heading(doc, "سجل الأدلة", 2)
    add_grid(doc, ["#", "الدليل", "ما يثبته", "تاريخه", "المالك", "مرجع الحفظ"], [
        [str(i), "[اسم موحد]", "[العلاقة بالنتيجة]", "[تاريخ]", "[الدور]", "[مسار داخلي]"] for i in range(1, 4)
    ], widths=[0.8, 4.3, 4.1, 2.4, 2.5, 3.3], font_size=7.8)
    add_approval_page(doc, ["يتضمن التقرير نتيجة وتحليلًا، لا وصف نشاط فقط.", "يمكن تتبع كل حكم إلى دليل محفوظ ومؤرخ."])
    return doc


def build_activity_impact() -> Document:
    title = "قالب تقرير نشاط وقياس أثر"
    doc = Document()
    set_document_defaults(doc, title)
    add_cover(doc, title, "توثيق نشاط يقيس الوصول والرضا والتعلم أو التغير، ويحفظ الخصوصية.", "activity-impact")
    doc.add_page_break()
    add_heading(doc, "تصميم النشاط", 1)
    add_fields_table(doc, [
        ("اسم النشاط", "[اسم واضح]"), ("نوع النشاط", "ورشة / لقاء / حملة / مبادرة"),
        ("الحاجة المثبتة", "[نتيجة/فجوة دفعت للنشاط]"), ("الفئة المستهدفة", "[وصف دون بيانات شخصية]"),
        ("الهدف القابل للقياس", "[نتيجة + قيمة + وقت]"), ("المالك", "[الدور/اللجنة]"),
        ("التاريخ والمكان", "[تاريخ] · [حضوري/افتراضي]"), ("الموارد والموافقات", "[موارد/مرجع موافقة]"),
    ])
    add_heading(doc, "خطة القياس", 2)
    add_grid(doc, ["مستوى القياس", "المؤشر", "خط الأساس", "المستهدف", "الأداة", "توقيت القياس"], [
        ["الوصول", "[عدد/نسبة]", "[قيمة]", "[قيمة]", "سجل حضور", "أثناء التنفيذ"],
        ["الرضا", "[متوسط/نسبة]", "[قيمة]", "[قيمة]", "استبانة", "بعد النشاط"],
        ["التعلم/الأثر", "[فرق/تطبيق]", "[قيمة]", "[قيمة]", "قبلي-بعدي/متابعة", "[موعد]"]
    ], widths=[2.6, 3.2, 2.5, 2.5, 3.6, 3.0], font_size=7.8)
    add_heading(doc, "قائمة جاهزية التنفيذ", 2)
    add_checklist(doc, ["الموافقة والموارد متاحة.", "أداة القياس مجربة ومفهومة.", "المشاركون أُبلغوا بالغرض والخصوصية.", "خطة الطوارئ وقناة المسؤولية محددتان."])
    doc.add_page_break()
    add_heading(doc, "التنفيذ والنتائج", 1)
    add_grid(doc, ["العنصر", "المخطط", "الفعلي", "الفرق", "التفسير", "الشاهد"], [
        ["المستفيدون", "[عدد]", "[عدد]", "[±]", "[سبب]", "[مرجع]"],
        ["المدة/الجلسات", "[قيمة]", "[قيمة]", "[±]", "[سبب]", "[مرجع]"],
        ["الرضا", "[قيمة]", "[قيمة]", "[±]", "[تفسير]", "[تحليل]"],
        ["الأثر", "[قيمة]", "[قيمة]", "[±]", "[تفسير]", "[أداة]"]
    ], widths=[2.6, 2.4, 2.4, 1.8, 5.0, 3.2], font_size=7.8)
    add_heading(doc, "تحليل الأثر والقيود", 2)
    add_fields_table(doc, [("أهم نتيجة", "[نتيجة مع رقم]"), ("ما الذي دعمها؟", "[عامل أو عنصر تصميم]"), ("ما الذي حد منها؟", "[قيد/تحيز/نقص استجابة]"), ("قرار التكرار أو التعديل", "[قرار مبرر]")])
    add_heading(doc, "توثيق يحفظ الخصوصية", 2)
    add_callout(doc, "قاعدة", "لا تُدرج أرقام اتصال أو قوائم أسماء أو صورًا قابلة للتعرف إلا عند الضرورة وبصلاحية مناسبة. استخدم إجماليات وروابط داخلية مقيدة.", tone="gold")
    add_approval_page(doc, ["أرفقت أداة القياس ونتيجتها المجمعة.", "قيس الأثر أو التعلم ولم يقتصر التقرير على عدد الحضور والصور."])
    return doc


def build_improvement_report() -> Document:
    title = "قالب تقرير معالجة ملاحظة وفرصة تحسين"
    doc = Document()
    set_document_defaults(doc, title)
    add_cover(doc, title, "سجل تحليلي من الملاحظة إلى إثبات فاعلية الإجراء وإقفالها.", "improvement-report")
    doc.add_page_break()
    add_heading(doc, "تعريف الملاحظة وتقييمها", 1)
    add_fields_table(doc, [
        ("رقم الملاحظة", "[IMPR-###]"), ("المصدر", "تقرير / مراجعة / محضر / مؤشر / استبانة"),
        ("النص الدقيق", "[الملاحظة دون إعادة تأويل]"), ("تاريخ الرصد", "[تاريخ]"),
        ("المعيار/المحك", "[مرجع إن وجد]"), ("المالك", "[الدور/اللجنة]"),
        ("درجة الأولوية", "حرجة / عالية / متوسطة / منخفضة"), ("مبرر الأولوية", "[أثر × احتمال × اتساع]"),
    ])
    add_heading(doc, "تحليل السبب الجذري", 2)
    add_grid(doc, ["السؤال", "التحليل", "الدليل المؤيد", "درجة الثقة"], [
        ["ما المشكلة القابلة للملاحظة؟", "[وصف]", "[مصدر]", "عالٍ/متوسط/منخفض"],
        ["لماذا حدثت؟", "[سبب 1]", "[مصدر]", "عالٍ/متوسط/منخفض"],
        ["لماذا استمر السبب؟", "[سبب 2]", "[مصدر]", "عالٍ/متوسط/منخفض"],
        ["ما السبب الجذري المرجح؟", "[استنتاج]", "[حزمة أدلة]", "عالٍ/متوسط/منخفض"],
    ], widths=[4.4, 5.2, 4.3, 3.5], font_size=7.9)
    add_callout(doc, "ضابط", "إذا لم يكفِ الدليل لتحديد سبب جذري، سجّل فرضية واختبار تحقق بدل تقديمها كحقيقة.", tone="gold")
    doc.add_page_break()
    add_heading(doc, "خطة المعالجة والتحقق", 1)
    add_grid(doc, ["#", "الإجراء", "النوع", "المسؤول", "الموعد", "مؤشر الفاعلية", "الدليل"], [
        [str(i), "[إجراء محدد]", "تصحيحي/وقائي", "[الدور]", "[تاريخ]", "[مؤشر + هدف]", "[متوقع]"] for i in range(1, 5)
    ], widths=[0.7, 3.6, 2.6, 2.4, 2.2, 3.5, 2.4], font_size=7.35)
    add_heading(doc, "نتيجة التحقق من الفاعلية", 2)
    add_grid(doc, ["المؤشر", "قبل", "المستهدف", "بعد", "المصدر", "الحكم"], [
        ["[مؤشر]", "[قيمة]", "[قيمة]", "[قيمة]", "[مصدر]", "فعال/جزئي/غير فعال"] for _ in range(3)
    ], widths=[3.8, 2.2, 2.5, 2.2, 3.3, 3.4], font_size=7.8)
    add_heading(doc, "قرار الإقفال", 2)
    add_fields_table(doc, [("القرار", "إقفال / تمديد / إعادة فتح"), ("مبرر القرار", "[استنادًا إلى نتيجة القياس]"), ("الإجراء المتبقي", "[إن وجد]"), ("موعد المراجعة اللاحقة", "[تاريخ/لا ينطبق]")])
    add_approval_page(doc, ["لم يُغلق الإجراء بمجرد التنفيذ؛ تحققت فاعليته بمؤشر.", "فُصلت الحقائق عن الفرضيات وذُكرت درجة الثقة."])
    return doc


def build_final_report() -> Document:
    title = "قالب التقرير الختامي للجنة"
    doc = Document()
    set_document_defaults(doc, title)
    add_cover(doc, title, "تجميع تنفيذي لأعمال اللجنة ونتائجها وتحسيناتها خلال دورة العمل.", "final-report")
    doc.add_page_break()
    add_heading(doc, "الملخص التنفيذي والحوكمة", 1)
    add_fields_table(doc, [
        ("اسم اللجنة", "[الاسم الرسمي]"), ("مرجع التشكيل", "[قرار/تاريخ]"),
        ("فترة التقرير", "[من] إلى [إلى]"), ("الرئيس والمقرر", "[الأدوار]"),
        ("عدد الاجتماعات", "[عدد مدعوم بسجل]"), ("عدد المهام", "[مخطط / منجز / غير منجز]"),
        ("أبرز نتيجة", "[نتيجة كمية أو نوعية موثقة]"), ("أبرز تحدٍ", "[تحدٍ وتأثيره]"),
    ])
    add_heading(doc, "مؤشرات موجزة", 2)
    add_grid(doc, ["المؤشر", "طريقة الحساب", "الخط الأساس", "المستهدف", "النتيجة", "الاتجاه"], [
        ["[اسم المؤشر]", "[بسط/مقام ×100]", "[قيمة]", "[قيمة]", "[قيمة]", "تحسن/ثبات/تراجع"] for _ in range(4)
    ], widths=[3.6, 4.0, 2.3, 2.3, 2.3, 2.9], font_size=7.7)
    add_callout(doc, "قاعدة الملخص", "اكتب أهم استنتاج وقرار أولًا، ثم دع التفاصيل والسجلات تثبته. لا تستخدم نسب إنجاز غير مدعومة بسجل تسليم.", tone="gold")
    doc.add_page_break()
    add_heading(doc, "نتائج الأعمال وتحليل الأداء", 1)
    add_grid(doc, ["#", "المهمة/المخرج", "الحالة", "النتيجة", "التحليل", "الدليل"], [
        [str(i), "[مهمة محددة]", "مكتمل/جزئي/مرحل", "[نتيجة]", "[سبب/أثر]", "[مرجع]"] for i in range(1, 6)
    ], widths=[0.7, 3.7, 2.4, 3.6, 4.1, 2.9], font_size=7.45)
    add_heading(doc, "القرارات والإجراءات المستمرة", 2)
    add_grid(doc, ["#", "القرار/الإجراء", "المسؤول", "الموعد", "الحالة", "دليل الإقفال"], [
        [str(i), "[صياغة واضحة]", "[الدور]", "[تاريخ]", "مفتوح/مغلق", "[مرجع]"] for i in range(1, 4)
    ], widths=[0.8, 4.6, 3.0, 2.7, 2.8, 3.5], font_size=7.8)
    add_heading(doc, "التحسين للدورة التالية", 2)
    add_grid(doc, ["الأولوية", "الدليل", "التحسين المقترح", "المستهدف", "المالك", "الموعد"], [
        ["عالية/متوسطة", "[نتيجة التقرير]", "[تحسين]", "[قيمة/نتيجة]", "[الدور]", "[تاريخ]"] for _ in range(3)
    ], widths=[2.3, 3.4, 4.7, 2.9, 2.2, 1.9], font_size=7.8)
    add_approval_page(doc, ["تتطابق أعداد المهام والاجتماعات مع السجلات المصدرية.", "لا تُعرض حالة تسليم فعلية ما لم تثبت من SharePoint أو سجل معتمد."])
    return doc


BUILDERS = {
    "قالب_محضر_اجتماع_لجنة.docx": build_minutes,
    "قالب_خطة_تشغيلية_للجنة.docx": build_operational_plan,
    "قالب_تقرير_إنجاز_وتحليل_مهمة.docx": build_completion_report,
    "قالب_تقرير_نشاط_وقياس_أثر.docx": build_activity_impact,
    "قالب_تقرير_معالجة_ملاحظة_وفرصة_تحسين.docx": build_improvement_report,
    "قالب_التقرير_الختامي_للجنة.docx": build_final_report,
}


def main() -> None:
    for filename, builder in BUILDERS.items():
        document = builder()
        path = OUTPUT_DIR / filename
        document.save(path)
        print(path)


if __name__ == "__main__":
    main()
